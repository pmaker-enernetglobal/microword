#! /usr/bin/env python
#
# microword.py - converts a very simple plain text file into microsoftword
#

'''
Converts a small markup language from plain text to microft word and saves it 
to output.docx
'''

# TODO: a air bit of tidying up

from docx import Document
import sys

doc = Document()
table = None
para = None
for line in sys.stdin:
    line.lstrip()
    if line[0] == '*':
        n = 0
        while line[0] == '*':
            n += 1
            line = line[1:]
        doc.add_heading(line, level=n)
        doc.add_paragraph('')
        table = None
    elif line[0] == '|':
        cn = line.count('|') - 1
        if table == None:
            table = doc.add_table(rows=1, cols=cn, style='Light Grid')
            table.autofit = True
            columns = cn
            row = 0
        else:
            if cn != columns:
                print(line, 'wrong # of columns')
        table.add_row()
        for c, f in enumerate(line.split('|')[1:-1]):
            table.cell(row, c).text = f
        row += 1
    elif line == '\n':
        table = None
        if para != None:
            doc.add_paragraph(para)
            para = None
        else:
            pass
    else:
        table = None
        if para == None:
            para = ''
        para += line

if para != None:
    doc.add_paragraph(para)
        
doc.save('output.docx')
